import os
import re
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from app.core.logging import logger

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None


@dataclass
class ParsedPage:
    page_number: int
    raw_text: str
    word_count: int


@dataclass
class ParsedDocumentMetadata:
    title: Optional[str] = None
    author: Optional[str] = None
    page_count: int = 1
    total_word_count: int = 0
    file_type: str = "txt"
    extra_metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    filename: str
    metadata: ParsedDocumentMetadata
    pages: List[ParsedPage]


class BaseParser:
    """Abstract Base Class for multi-format document parsers."""

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        raise NotImplementedError


class PDFParser(BaseParser):
    """
    PDF parser preserving page numbers and extracting PDF metadata.
    """

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        pages: List[ParsedPage] = []
        title: Optional[str] = None
        author: Optional[str] = None
        extra_meta: Dict[str, Any] = {}

        if pypdf:
            try:
                reader = pypdf.PdfReader(file_path)
                pdf_meta = reader.metadata or {}
                if pdf_meta.title:
                    title = str(pdf_meta.title).strip()
                if pdf_meta.author:
                    author = str(pdf_meta.author).strip()
                if pdf_meta.creator:
                    extra_meta["creator"] = str(pdf_meta.creator)

                for idx, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    clean_page_text = page_text.strip()
                    word_count = len(clean_page_text.split()) if clean_page_text else 0
                    pages.append(
                        ParsedPage(
                            page_number=idx + 1,
                            raw_text=clean_page_text,
                            word_count=word_count,
                        )
                    )
            except Exception as e:
                logger.error(f"PyPDF parsing failed for {filename}: {e}. Falling back to text fallback.")

        # Fallback if pages empty or pypdf not available
        if not pages:
            pages = [ParsedPage(page_number=1, raw_text=f"[PDF parsing error for {filename}]", word_count=5)]

        # Fallback title heuristic from first page text line if missing
        if not title and pages and pages[0].raw_text:
            first_line = pages[0].raw_text.split("\n")[0].strip()
            if 3 < len(first_line) < 120:
                title = first_line

        if not title:
            title = os.path.splitext(filename)[0].replace("_", " ").title()

        total_words = sum(p.word_count for p in pages)
        meta = ParsedDocumentMetadata(
            title=title,
            author=author,
            page_count=len(pages),
            total_word_count=total_words,
            file_type="pdf",
            extra_metadata=extra_meta,
        )
        return ParsedDocument(filename=filename, metadata=meta, pages=pages)


class DOCXParser(BaseParser):
    """
    DOCX parser reading paragraphs and section structures.
    """

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        pages: List[ParsedPage] = []
        title: Optional[str] = None
        author: Optional[str] = None

        if docx:
            try:
                doc = docx.Document(file_path)
                core_props = doc.core_properties
                if core_props.title:
                    title = core_props.title.strip()
                if core_props.author:
                    author = core_props.author.strip()

                full_text = []
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        full_text.append(text)

                combined_text = "\n\n".join(full_text)
                words = combined_text.split()
                total_words = len(words)

                # Segment into ~500 word pages for chunking & pagination preservation
                words_per_page = 500
                if total_words == 0:
                    pages.append(ParsedPage(page_number=1, raw_text="", word_count=0))
                else:
                    current_page_words = []
                    page_num = 1
                    for word in words:
                        current_page_words.append(word)
                        if len(current_page_words) >= words_per_page:
                            p_text = " ".join(current_page_words)
                            pages.append(
                                ParsedPage(
                                    page_number=page_num,
                                    raw_text=p_text,
                                    word_count=len(current_page_words),
                                )
                            )
                            page_num += 1
                            current_page_words = []
                    if current_page_words:
                        p_text = " ".join(current_page_words)
                        pages.append(
                            ParsedPage(
                                page_number=page_num,
                                raw_text=p_text,
                                word_count=len(current_page_words),
                            )
                        )
            except Exception as e:
                logger.error(f"DOCX parsing failed for {filename}: {e}")

        if not pages:
            pages = [ParsedPage(page_number=1, raw_text="", word_count=0)]

        if not title:
            title = os.path.splitext(filename)[0].replace("_", " ").title()

        meta = ParsedDocumentMetadata(
            title=title,
            author=author,
            page_count=len(pages),
            total_word_count=sum(p.word_count for p in pages),
            file_type="docx",
        )
        return ParsedDocument(filename=filename, metadata=meta, pages=pages)


class TXTParser(BaseParser):
    """
    Plain text document parser.
    """

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        except Exception as e:
            logger.error(f"TXT read failed: {e}")
            content = ""

        words = content.split()
        words_per_page = 500
        pages: List[ParsedPage] = []

        if not words:
            pages.append(ParsedPage(page_number=1, raw_text="", word_count=0))
        else:
            page_num = 1
            for i in range(0, len(words), words_per_page):
                chunk_words = words[i : i + words_per_page]
                p_text = " ".join(chunk_words)
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        raw_text=p_text,
                        word_count=len(chunk_words),
                    )
                )
                page_num += 1

        title = os.path.splitext(filename)[0].replace("_", " ").title()
        meta = ParsedDocumentMetadata(
            title=title,
            page_count=len(pages),
            total_word_count=len(words),
            file_type="txt",
        )
        return ParsedDocument(filename=filename, metadata=meta, pages=pages)


class CSVParser(BaseParser):
    """
    CSV dataset parser extracting tabular headers and text representation.
    """

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        extra_meta: Dict[str, Any] = {}
        pages: List[ParsedPage] = []
        raw_text_summary = ""

        if pd is not None:
            try:
                df = pd.read_csv(file_path)
                extra_meta["columns"] = list(df.columns)
                extra_meta["row_count"] = len(df)
                extra_meta["column_count"] = len(df.columns)

                # Create descriptive textual representation of dataset for RAG / NLP
                summary_lines = [
                    f"CSV Dataset: {filename}",
                    f"Total Rows: {len(df)}, Total Columns: {len(df.columns)}",
                    f"Columns: {', '.join(df.columns)}",
                    "\nDataset Sample (Top 5 rows):",
                    df.head(5).to_string(),
                ]
                raw_text_summary = "\n".join(summary_lines)
            except Exception as e:
                logger.error(f"CSV read failed: {e}")

        if not raw_text_summary:
            raw_text_summary = f"CSV dataset file: {filename}"

        words = raw_text_summary.split()
        pages.append(ParsedPage(page_number=1, raw_text=raw_text_summary, word_count=len(words)))

        title = f"Dataset: {os.path.splitext(filename)[0].replace('_', ' ').title()}"
        meta = ParsedDocumentMetadata(
            title=title,
            page_count=1,
            total_word_count=len(words),
            file_type="csv",
            extra_metadata=extra_meta,
        )
        return ParsedDocument(filename=filename, metadata=meta, pages=pages)


class DocumentParserFactory:
    """
    Factory creating parser instances based on file extension.
    """

    @staticmethod
    def get_parser(filename: str) -> BaseParser:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return PDFParser()
        elif ext in [".docx", ".doc"]:
            return DOCXParser()
        elif ext == ".csv":
            return CSVParser()
        elif ext in [".txt", ".md", ".json"]:
            return TXTParser()
        else:
            return TXTParser()
